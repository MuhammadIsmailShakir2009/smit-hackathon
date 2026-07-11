"""
resume_parser.py
-----------------
Yeh module "PDF modality" ko handle karta hai. Resume files (PDF, DOCX, TXT)
se raw text extract karne ka kaam yahan hota hai. Yeh multimodal pipeline
ka pehla step hai jahan unstructured document data ko text mein convert
kiya jata hai taake aage NLP processing ho sake.
"""

import io
import pdfplumber
import docx


def parse_pdf(file_bytes: bytes) -> str:
    """PDF file ke bytes se pura text nikalta hai (PDF modality)."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """DOCX file ke bytes se text nikalta hai."""
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

    # Tables ke andar bhi text ho sakta hai isliye unko bhi parse karte hain
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def parse_txt(file_bytes: bytes) -> str:
    """Plain text file ko decode karta hai."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def parse_resume(file_name: str, file_bytes: bytes) -> str:
    """
    File extension check kar ke sahi parser call karta hai.
    Yeh function resume_parser module ka main entry point hai.
    """
    lower_name = file_name.lower()

    if lower_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return parse_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        return parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format for '{file_name}'. Sirf PDF, DOCX aur TXT supported hain."
        )
